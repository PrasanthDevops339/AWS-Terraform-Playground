"""
Script to generate the Suspended OU Logic Word document.
Run: python generate_suspended_ou_doc.py
Output: Suspended_OU_Logic.docx in the same directory
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    from lxml import etree
    shading = cell._element.get_or_add_tcPr()
    shading_elem = etree.SubElement(shading, qn('w:shd'))
    shading_elem.set(qn('w:fill'), color_hex)
    shading_elem.set(qn('w:val'), 'clear')


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2E74B5')

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            # Alternate row shading
            if r % 2 == 1:
                set_cell_shading(cell, 'DEEAF6')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def main():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Suspended OU Exclusion Logic', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AWS Config Aggregator - config_aggregator.py')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    # --- Section 1: Overview ---
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'The config_aggregator.py script queries AWS Config for non-compliant resources '
        'across all accounts in an AWS Organization. Accounts placed under the Suspended OU '
        'are decommissioned or inactive and must be excluded from compliance reporting.'
    )
    doc.add_paragraph(
        'This document explains the suspended OU exclusion logic, including how the cache works, '
        'the two-layer defense mechanism, and how different accounts are handled.'
    )

    # --- Section 2: What is a Suspended OU? ---
    doc.add_heading('2. What Is a Suspended OU?', level=1)
    doc.add_paragraph(
        'In AWS Organizations, an Organizational Unit (OU) is a logical grouping of accounts. '
        'A Suspended OU is a dedicated OU where decommissioned, inactive, or retired accounts '
        'are moved. These accounts are excluded from compliance reporting because:'
    )
    bullet_items = [
        'They are no longer actively managed.',
        'Resources in those accounts are scheduled for cleanup or deletion.',
        'Reporting on them creates noise and wastes processing time.',
    ]
    for item in bullet_items:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Hardcoded OU ID: ')
    run.bold = True
    p.add_run('SUSPENDED_OU_ID = \'ou-susp345jkl\'')

    # --- Section 3: Sample Accounts ---
    doc.add_heading('3. Sample Accounts', level=1)
    doc.add_paragraph(
        'The following sample accounts illustrate how the exclusion logic applies to '
        'different scenarios:'
    )

    headers = ['Account ID', 'Account Name', 'OU Location', 'Suspended?',
               '1.0 (Legacy)?', 'Final Result']
    rows = [
        ['111122223333', 'prod-app-account', 'Production OU', 'No', 'No', 'INCLUDED in report'],
        ['222233334444', 'dev-test-account', 'Development OU', 'No', 'No', 'INCLUDED in report'],
        ['333344445555', 'legacy-finance', 'Production OU', 'No', 'Yes', 'Excluded (1.0 legacy)'],
        ['444455556666', 'old-project', 'Suspended OU', 'YES', 'No', 'Excluded (Suspended)'],
        ['555566667777', 'decommissioned-app', 'Suspended OU', 'YES', 'No', 'Excluded (Suspended)'],
        ['666677778888', 'retired-legacy-app', 'Suspended OU', 'YES', 'Yes',
         'Excluded (Suspended - checked first)'],
    ]
    add_styled_table(doc, headers, rows)

    # --- Section 4: How the Exclusion Works ---
    doc.add_heading('4. How the Suspended OU Exclusion Works', level=1)

    # Step 1
    doc.add_heading('Step 1: Build a Cache of Suspended Account IDs', level=2)
    doc.add_paragraph(
        'When the script first needs to check whether an account is suspended, it calls '
        'get_suspended_account_ids(). This function:'
    )
    steps = [
        'Calls the AWS Organizations API (list_accounts_for_parent) with the Suspended OU ID.',
        'Collects all account IDs directly under that OU (nested child OUs are NOT traversed).',
        'Stores the result in a Python set for O(1) lookup performance.',
        'Caches the set with a timestamp so it is not re-fetched on every call.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Cache Configuration:')
    run.bold = True

    cache_headers = ['Variable', 'Default', 'Description']
    cache_rows = [
        ['SUSPENDED_OU_CACHE_TTL_ENABLED', 'true',
         'Whether the cache expires after a TTL period'],
        ['SUSPENDED_OU_CACHE_TTL_SECONDS', '1800',
         'Seconds before cache expires (30 minutes)'],
    ]
    add_styled_table(doc, cache_headers, cache_rows)

    doc.add_paragraph()
    doc.add_paragraph(
        'If TTL is disabled (SUSPENDED_OU_CACHE_TTL_ENABLED=false), the cache persists for '
        'the entire lifetime of the process.'
    )

    # Step 2
    doc.add_heading('Step 2: Per-Resource Check (Layer 1)', level=2)
    doc.add_paragraph(
        'During the main processing loop, after fetching non-compliant resources from AWS Config, '
        'the script iterates over each resource. Before any other processing, it checks whether '
        'the resource\'s account is in the Suspended OU.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Code (Line 394-396 in config_aggregator.py):')
    run.bold = True
    run.font.size = Pt(9)

    code_block = doc.add_paragraph()
    code_run = code_block.add_run(
        'if is_account_in_suspended_ou(account_id):\n'
        '    logger.warning(f"[SUSPENDED OU] Skipping account: {account_id} ({account_name})")\n'
        '    continue  # Skip this resource entirely'
    )
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    doc.add_paragraph(
        'This is the primary filter. If the account is suspended, the resource is skipped '
        'immediately and never written to the in-memory CSV buffer.'
    )

    # Step 3
    doc.add_heading('Step 3: CSV Upload Safety Check (Layer 2)', level=2)
    doc.add_paragraph(
        'After all resources are processed and grouped by account, the script performs a second '
        'check before uploading each account\'s CSV to S3.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Code (Lines 461-463 in config_aggregator.py):')
    run.bold = True
    run.font.size = Pt(9)

    code_block = doc.add_paragraph()
    code_run = code_block.add_run(
        'if is_account_in_suspended_ou(group_account_id):\n'
        '    logger.warning(f"[SUSPENDED OU] Skipping CSV creation for suspended account: '
        '{group_account_id}")\n'
        '    continue  # Do not create or upload CSV'
    )
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    doc.add_paragraph(
        'This is a safety net. Even if a suspended account\'s resource somehow made it into the '
        'CSV buffer (e.g., due to a race condition or caching edge case), this check prevents '
        'the CSV file from being created and uploaded to S3.'
    )

    # --- Section 5: Two-Layer Defense Summary ---
    doc.add_heading('5. Two-Layer Defense Summary', level=1)

    defense_headers = ['Layer', 'Where in Code', 'What It Prevents', 'Why It Exists']
    defense_rows = [
        ['Layer 1', 'Resource iteration (line 394)',
         'Resources written to CSV buffer', 'Primary filter - fast, efficient'],
        ['Layer 2', 'CSV upload (line 464)',
         'CSV file creation and S3 upload', 'Safety net - defense in depth'],
    ]
    add_styled_table(doc, defense_headers, defense_rows)

    # --- Section 6: Decision Flowchart ---
    doc.add_heading('6. Decision Flowchart', level=1)
    doc.add_paragraph(
        'For each non-compliant resource returned by AWS Config, the following decision '
        'process is applied:'
    )

    flow_steps = [
        ('Step A', 'Get the account_id and account_name for the resource.'),
        ('Step B', 'CHECK: Is account in Suspended OU? '
                   'If YES -> SKIP resource (log warning). If NO -> continue.'),
        ('Step C', 'Retrieve config rule annotations for the resource.'),
        ('Step D', 'CHECK: Is account 1.0 (legacy) in DynamoDB? '
                   'If YES -> SKIP resource (log: 1.0 account). If NO -> continue.'),
        ('Step E', 'CHECK: Does the resource have matching annotations? '
                   'If YES -> INCLUDE in CSV report and upload to S3. '
                   'If NO -> SKIP resource (no matching annotations).'),
    ]
    for label, desc in flow_steps:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        p.add_run(desc)

    # --- Section 7: Walkthrough with Sample Accounts ---
    doc.add_heading('7. Walkthrough: Sample Account Processing', level=1)

    # Account 444455556666
    doc.add_heading('Account 444455556666 (old-project) - Suspended', level=2)
    walkthrough_1 = [
        'Resource vol-0abc123def456 found as NON_COMPLIANT.',
        'Account name resolved: old-project.',
        'Suspended OU check: 444455556666 IS in the cached suspended set.',
        'Result: Resource is SKIPPED.',
        'Log output: [SUSPENDED OU] Skipping account: 444455556666 (old-project)',
        'No CSV is written, no S3 upload occurs.',
    ]
    for i, step in enumerate(walkthrough_1, 1):
        doc.add_paragraph(f'{i}. {step}')

    # Account 333344445555
    doc.add_heading('Account 333344445555 (legacy-finance) - 1.0 Legacy', level=2)
    walkthrough_2 = [
        'Resource vol-0xyz789abc123 found as NON_COMPLIANT.',
        'Account name resolved: legacy-finance.',
        'Suspended OU check: 333344445555 is NOT in the suspended set. Continue.',
        'Config rule annotations retrieved.',
        '1.0 check: DynamoDB query finds legacy-finance in the version table.',
        'Result: Resource is SKIPPED.',
        'Log output: Account 333344445555 & legacy-finance is 1.0 - skipping',
    ]
    for i, step in enumerate(walkthrough_2, 1):
        doc.add_paragraph(f'{i}. {step}')

    # Account 111122223333
    doc.add_heading('Account 111122223333 (prod-app-account) - Included', level=2)
    walkthrough_3 = [
        'Resource vol-0def456ghi789 found as NON_COMPLIANT.',
        'Account name resolved: prod-app-account.',
        'Suspended OU check: 111122223333 is NOT in the suspended set. Continue.',
        'Config rule annotations retrieved: ["Unencrypted volume"].',
        '1.0 check: DynamoDB query does NOT find prod-app-account. Account is 2.0.',
        'Annotations are present.',
        'Result: Resource is INCLUDED in CSV and uploaded to S3.',
        'Log output: CSV saved to s3://my-bucket/config-reports/111122223333-prod-app-account_1.csv',
    ]
    for i, step in enumerate(walkthrough_3, 1):
        doc.add_paragraph(f'{i}. {step}')

    # Account 666677778888
    doc.add_heading(
        'Account 666677778888 (retired-legacy-app) - Suspended AND 1.0', level=2
    )
    walkthrough_4 = [
        'Resource found as NON_COMPLIANT.',
        'Account name resolved: retired-legacy-app.',
        'Suspended OU check: 666677778888 IS in the cached suspended set.',
        'Result: Resource is SKIPPED at the Suspended OU check. '
        'The DynamoDB (1.0) check is NEVER reached.',
        'Log output: [SUSPENDED OU] Skipping account: 666677778888 (retired-legacy-app)',
    ]
    for i, step in enumerate(walkthrough_4, 1):
        doc.add_paragraph(f'{i}. {step}')

    # --- Section 8: Precedence / Overlap ---
    doc.add_heading('8. What Happens When an Account Is in Both Suspended OU AND DynamoDB?', level=1)
    doc.add_paragraph(
        'The Suspended OU check runs FIRST (line 394). If an account is in the Suspended OU, '
        'the continue statement immediately skips to the next resource. The DynamoDB 1.0 check '
        '(line 416) is never reached.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Why this order is intentional:')
    run.bold = True
    doc.add_paragraph(
        '1. Suspended OU check = O(1) in-memory set lookup (cached, no API call).',
    )
    doc.add_paragraph(
        '2. DynamoDB check = network call to DynamoDB (slower, costs money).',
    )
    doc.add_paragraph(
        'By checking the cached Suspended OU first, the script avoids unnecessary DynamoDB '
        'queries for suspended accounts, saving time and reducing costs.'
    )

    doc.add_paragraph()
    scenario_headers = ['Scenario', 'In Suspended OU?', 'In DynamoDB (1.0)?',
                        'What Happens', 'Log Message']
    scenario_rows = [
        ['Both flags', 'Yes', 'Yes',
         'Skipped at Suspended OU (DynamoDB never queried)',
         '[SUSPENDED OU] Skipping account...'],
        ['Only Suspended', 'Yes', 'No',
         'Skipped at Suspended OU check',
         '[SUSPENDED OU] Skipping account...'],
        ['Only 1.0', 'No', 'Yes',
         'Skipped at DynamoDB check',
         'Account ... is 1.0 - skipping'],
        ['Neither (with annotations)', 'No', 'No',
         'INCLUDED in report',
         '(no skip log, writes to CSV)'],
        ['Neither (no annotations)', 'No', 'No',
         'Skipped (no matching annotations)',
         'Account ... - skipped (no matching annotations)'],
    ]
    add_styled_table(doc, scenario_headers, scenario_rows)

    # --- Section 9: Key Functions ---
    doc.add_heading('9. Key Functions Reference', level=1)

    func_headers = ['Function', 'Purpose', 'Returns']
    func_rows = [
        ['_list_accounts_for_parent()', 'Lists accounts directly under a given OU',
         'List of account ID strings'],
        ['get_suspended_account_ids()', 'Fetches and caches suspended account IDs',
         'Set of account ID strings'],
        ['is_account_in_suspended_ou()', 'Checks if a single account is in the Suspended OU',
         'True or False'],
    ]
    add_styled_table(doc, func_headers, func_rows)

    # --- Section 10: Important Notes ---
    doc.add_heading('10. Important Notes', level=1)

    notes = [
        ('Only direct children', 'The script only checks accounts directly under the '
         'Suspended OU. Accounts in nested child OUs within the Suspended OU are NOT excluded.'),
        ('Error handling', 'If the Organizations API call fails, the function returns an empty '
         'set, so no valid accounts are accidentally skipped.'),
        ('Execution order', 'The Suspended OU check runs before the DynamoDB 1.0 check. This '
         'is intentional -- the in-memory set lookup is faster than a DynamoDB network call.'),
        ('Required IAM permission', 'The script needs organizations:ListAccountsForParent to '
         'fetch accounts from the Suspended OU.'),
    ]
    for heading, detail in notes:
        p = doc.add_paragraph()
        run = p.add_run(f'{heading}: ')
        run.bold = True
        p.add_run(detail)

    # --- Section 11: Log Output Examples ---
    doc.add_heading('11. Sample Log Output', level=1)

    doc.add_heading('Suspended OU Cache Initialization', level=2)
    log_block = doc.add_paragraph()
    log_run = log_block.add_run(
        '[SUSPENDED OU] Cached 3 suspended accounts\n'
        '[SUSPENDED OU] Account IDs: 444455556666, 555566667777, 666677778888'
    )
    log_run.font.name = 'Courier New'
    log_run.font.size = Pt(9)

    doc.add_heading('Account Processing', level=2)
    log_block = doc.add_paragraph()
    log_run = log_block.add_run(
        '# Suspended account skipped\n'
        '[SUSPENDED OU] Skipping account: 444455556666 (old-project)\n\n'
        '# 1.0 account skipped\n'
        'Account 333344445555 & legacy-finance is 1.0 - skipping\n\n'
        '# 2.0 account included (no skip log, writes to CSV)\n'
    )
    log_run.font.name = 'Courier New'
    log_run.font.size = Pt(9)

    doc.add_heading('S3 Upload', level=2)
    log_block = doc.add_paragraph()
    log_run = log_block.add_run(
        '# Layer 2 safety net\n'
        '[SUSPENDED OU] Skipping CSV creation for suspended account: 444455556666\n\n'
        '# Successful uploads\n'
        'CSV saved to s3://my-bucket/config-reports/111122223333-prod-app-account_1.csv\n'
        'CSV saved to s3://my-bucket/config-reports/222233334444-dev-test-account_1.csv'
    )
    log_run.font.name = 'Courier New'
    log_run.font.size = Pt(9)

    # Save
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               'Suspended_OU_Logic.docx')
    doc.save(output_path)
    print(f'Word document saved to: {output_path}')


if __name__ == '__main__':
    main()
